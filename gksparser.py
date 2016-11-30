from openpyxl import load_workbook
from docx import Document
from json import dumps
from bs4 import BeautifulSoup

import requests

def test():
	# wb = load_workbook('data/vrp98-14.xlsx')
	# years = get_years(wb.get_sheet_by_name('Лист1'))
	# x = get_region_data(wb.get_sheet_by_name('Лист1'), years)
	x = get_region_salary()
	print(dumps(x, ensure_ascii=0, indent=2))


def get_region_vrp(_file = 'data/vrp98-14.xlsx'):
	""" ВРП по регионам по годам ПЕРЕПИСАТЬ"""

	def get_years(worksheet):
		""" возвращает года из листа отчета """
		res = dict()

		for x in [chr(x) for x in range(66,91)]:
			try:
				year = worksheet['%s6' % (x)].value.strip(' г.')
				res.update({x:year})
			except AttributeError:
				break
		return res

	worksheet = load_workbook(_file).get_sheet_by_name('Лист1')
	years = get_years(worksheet)
	_min = 9
	_max = int(str(worksheet.dimensions).split(':')[1][1:])

	res = dict()

	for row in range(_min, _max):
		region = worksheet['A%s' % (row)]
		if region.fill.bgColor.rgb == '00000000':
			resres = dict()
			for col in years.keys():
				_date = '31.12.%s' % (years[col])
				try:
					n1   = int(worksheet['%s%s' % (col,row)].value * 1000000)
				except ValueError:
					n1   = 0
				except TypeError:					
					n1   = 0
				resres.update({_date:n1})
			res.update({str(region.value).strip():resres})
	return res


def get_region_salary(_file = 'data/t4.xlsx'):
	""" Среднемесячная номинальная начисленная заработная плата работников по полному кругу организаций в целом по экономике по субъектам Российской Федерации за 2000-2015гг. """

	def get_years(worksheet):
		""" возвращает года из листа отчета """
		res = dict()

		for x in [chr(x) for x in range(66,91)]:
			year = worksheet['%s5' % (x)].value
			if year != None: res.update({x:year})
		return res

	wb = load_workbook(_file)
	worksheet = wb.get_sheet_by_name('Лист1')
	_min = 9
	_max = 104

	res = dict()
	years = get_years(worksheet)

	for row in range(_min,_max):
		region = worksheet['A%s' % (row)]
		if region.fill.bgColor.rgb == '00000000':
			resres = dict()
			for col in years.keys():
				_date = '31.12.%s' % (years[col])
				try:
					n1   = int(worksheet['%s%s' % (col,row)].value)
				except ValueError:
					n1   = 0
				except TypeError:					
					n1   = 0
				resres.update({_date:n1})
			res.update({str(region.value).strip():resres})
	return res

def get_region_income(_file = 'data/R_04-1.docx'):
	""" Среднедушевые Денежные Доходы Населения """
	doc = Document(_file)
	T = [doc.tables[2], doc.tables[3]]
	res = dict()
	n = 0
	for t in T:
		if n == 0:
			_min, _max = 1, 10
			n = 1
		else:
			_min, _max = 1, 10

		for row in range(_min, _max):
			try:
				_ = t.cell(row,0).text.upper().index("федера".upper())
			except ValueError:
				resres = dict()
				for col in range(1, 6):
					year = t.cell(0,col).text
					_date = '31.12.%s' % (year)
					try:
						n1 = int(t.cell(row,col).text)
					except Exception:
						n1 = 0
					resres.update({_date:n1})
				res.update({str(t.cell(row,0).text).strip():resres})
	return(res)

def get_region_rasxod(_file = 'data/R_04-1.docx'):
	""" Среднедушевые Денежные расходы Населения """
	doc = Document(_file)
	T = [doc.tables[36], doc.tables[37]]
	res = dict()
	n = 0
	for t in T:
		if n == 0:
			_min, _max = 1, 10
			n = 1
		else:
			_min, _max = 1, 10
			
		for row in range(_min, _max):
			try:
				_ = t.cell(row,0).text.upper().index("федера".upper())
			except ValueError:
				resres = dict()
				for col in range(1, 6):
					year = t.cell(0,col).text
					_date = '31.12.%s' % (year)
					try:
						n1 = int(t.cell(row,col).text)
					except Exception:
						n1 = 0
					resres.update({_date:n1})
				res.update({str(t.cell(row,0).text).strip():resres})
	return(res)

def get_region_population(url = 'http://www.gks.ru/bgd/regl/B14_14p/IssWWW.exe/Stg/d01/02-01.htm'):
	""" ЧИСЛЕННОСТЬ НАСЕЛЕНИЯ (оценка на конец года) """
	res = dict()
	soup = BeautifulSoup(requests.get(url).content)

	table = soup.find('table')
	rows = table('tr')
	years = [ int(p.text) for p in rows[0]('p') if p.text[0] == '2' ]

	res = {}
	for row in rows:
		try:
			if len(row('b')) == 0:
				resres = {}
				reg = row.td.font.p.text

				i = 0
				for td in row('td')[1:]:
					resres.update({'31.12.%s' % years[i]:int(td.text.strip('\n'))*1000})
					i += 1
				res.update({reg:resres})
		except AttributeError:
			pass
		except TypeError:
			pass
		except ValueError:
			pass
	return res

if __name__ == '__main__':
	# test()

	x = get_region_population()
	print(dumps(x, ensure_ascii=0, indent=2))
	exit(0)
	x = get_region_vrp()
	print(dumps(x, ensure_ascii=0, indent=2))
	x = get_region_salary()
	print(dumps(x, ensure_ascii=0, indent=2))
	x = get_region_income()
	print(dumps(x, ensure_ascii=0, indent=2))
	x = get_region_rasxod()
	print(dumps(x, ensure_ascii=0, indent=2))


	# http://www.gks.ru/wps/wcm/connect/rosstat_main/rosstat/ru/statistics/publications/catalog/doc_1138623506156    (Доходы - Р.4.2, Расходы  - Р.4.16)